r"""docx 脚本/プロットローダー.

検出ルール:
    - エピソード境界: 行頭が ``第\d+話`` または ``第[０-９]+話``
    - シーン境界: 行頭が ``〇`` または ``◯``（脚本のみ。診断時は参考程度）
    - 単一話 / プロット: エピソード境界が見つからない場合
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document

# 全角→半角の数字対応
_FW2HW = str.maketrans("０１２３４５６７８９", "0123456789")
_EPISODE_RE = re.compile(r"^第\s*([0-9０-９]+)\s*話")


@dataclass
class Episode:
    number: int
    title: str
    lines: list[str]

    @property
    def char_count(self) -> int:
        return sum(len(l) for l in self.lines)


@dataclass
class ScriptDocument:
    title: str
    header_lines: list[str]
    episodes: list[Episode]
    full_lines: list[str]

    @property
    def is_multi_episode(self) -> bool:
        return len(self.episodes) > 1

    @property
    def char_count(self) -> int:
        return sum(len(l) for l in self.full_lines)


def load_script(path: str | Path) -> ScriptDocument:
    """脚本/プロット .docx を読み込んでエピソード単位に分割する."""
    p = Path(path).expanduser()
    if not p.exists():
        raise FileNotFoundError(f"docx not found: {p}")

    doc = Document(str(p))
    raw = [para.text for para in doc.paragraphs]
    lines = [t.strip() for t in raw if t.strip()]

    episodes: list[Episode] = []
    current: Episode | None = None
    header: list[str] = []
    title = p.stem

    for line in lines:
        m = _EPISODE_RE.match(line)
        if m:
            if current is not None:
                episodes.append(current)
            num = int(m.group(1).translate(_FW2HW))
            # 「第1話」のあとにサブタイトルがあれば取り込む
            subtitle = line[m.end():].strip(" 　・\t-—")
            current = Episode(number=num, title=subtitle, lines=[])
        else:
            if current is None:
                header.append(line)
            else:
                current.lines.append(line)

    if current is not None:
        episodes.append(current)

    # エピソード境界が無ければ全文を 1 話として扱う
    if not episodes:
        episodes.append(Episode(number=1, title="", lines=lines.copy()))
        header = []

    return ScriptDocument(
        title=title,
        header_lines=header,
        episodes=episodes,
        full_lines=lines,
    )


def format_episode_index(doc: ScriptDocument) -> str:
    """エピソード一覧の人間可読サマリを返す."""
    rows = []
    for ep in doc.episodes:
        sub = f" {ep.title}" if ep.title else ""
        rows.append(
            f"  第{ep.number}話{sub}  ({len(ep.lines)} 行 / {ep.char_count:,} 字)"
        )
    return "\n".join(rows)
